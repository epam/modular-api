"""
Title page generator for DOCX documents
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.text.paragraph import Paragraph
from datetime import datetime
import logging

logger = logging.getLogger(__name__)


def add_title_page_to_docx(
        docx_path: str,
        title_text: str,
        secondary_title_text: str,
        date_text: str,
        version_text: str,
) -> None:
    """
    Add a title page to an existing DOCX document

    Args:
        docx_path: Path to the DOCX file
        title_text: Main title text (centered, large)
        secondary_title_text: Secondary title (bottom right)
        date_text: Date text (bottom right)
        version_text: Version text (bottom right)
    """
    try:
        doc = Document(docx_path)

        # Check if the first block in the document is a table
        first_element = doc.element.body[0]
        if first_element.tag == qn('w:tbl'):
            # Create a new paragraph element
            new_paragraph = OxmlElement('w:p')
            # Insert the new paragraph before the first table
            doc.element.body.insert(0, new_paragraph)

        first_para: Paragraph = doc.paragraphs[0]

        # Add spacing before title (12 empty lines)
        for _ in range(12):
            first_para.insert_paragraph_before()

        # Add title in the center of the page
        title = first_para.insert_paragraph_before()
        run = title.add_run(title_text)
        run.font.name = 'Arial'
        run.font.size = Pt(28)
        run.font.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add spacing after title (14 empty lines)
        for _ in range(14):
            first_para.insert_paragraph_before()

        # Add secondary title at the bottom of the page
        secondary_title = first_para.insert_paragraph_before()
        run = secondary_title.add_run(secondary_title_text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        secondary_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Add date
        date = first_para.insert_paragraph_before()
        run = date.add_run(date_text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Add spacing
        first_para.insert_paragraph_before()

        # Add version
        version = first_para.insert_paragraph_before()
        run = version.add_run(version_text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        version.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Add a page break to create a new page
        run = version.add_run()
        run.add_break(WD_BREAK.PAGE)

        # Save the document
        doc.save(docx_path)
        logger.info(f"Title page successfully added to: {docx_path}")
    except Exception as e:
        logger.error(f"Failed to add title page to {docx_path}: {e}")
        raise


def add_cli_title_page(
        docx_path: str,
        cli_name: str,
        document_code: str | None = None,
        version: str | None = None,
        date: str | None = None,
) -> None:
    """
    Add a title page for CLI documentation

    Args:
        docx_path: Path to the DOCX file
        cli_name: Name of the CLI tool
        document_code: Document code (e.g., 'SRE-DOC-01')
        version: Version string (e.g., 'Version 1.0')
        date: Date string (defaults to current month and year)
    """
    # Generate default values
    if date is None:
        date = datetime.now().strftime('%B %Y')

    if version is None:
        version = 'Version 1.0'

    if document_code is None:
        document_code = f"{cli_name.upper()}-DOC-01"

    # Format title text
    title_text = f"{cli_name}\nReference Guide"
    secondary_title = document_code

    add_title_page_to_docx(
        docx_path=docx_path,
        title_text=title_text,
        secondary_title_text=secondary_title,
        date_text=date,
        version_text=version,
    )


def add_custom_title_page(
        docx_path: str,
        main_title: str,
        subtitle: str | None = None,
        author: str | None = None,
        organization: str | None = None,
        version: str | None = None,
        date: str | None = None,
) -> None:
    """
    Add a customizable title page

    Args:
        docx_path: Path to the DOCX file
        main_title: Main title text
        subtitle: Subtitle text (optional)
        author: Author name (optional)
        organization: Organization name (optional)
        version: Version string (optional)
        date: Date string (defaults to current date)
    """
    if date is None:
        date = datetime.now().strftime('%B %d, %Y')

    # Build title text
    title_parts = [main_title]
    if subtitle:
        title_parts.append(subtitle)
    title_text = '\n'.join(title_parts)

    # Build secondary info
    secondary_parts = []
    if author:
        secondary_parts.append(f"Author: {author}")
    if organization:
        secondary_parts.append(organization)
    secondary_title = '\n'.join(secondary_parts) if secondary_parts else ''

    # Version text
    version_text = version if version else ''

    add_title_page_to_docx(
        docx_path=docx_path,
        title_text=title_text,
        secondary_title_text=secondary_title,
        date_text=date,
        version_text=version_text,
    )
