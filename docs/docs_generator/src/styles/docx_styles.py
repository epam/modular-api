"""
Custom styling for DOCX documents
"""

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging

CODE_BLOCK = 'Code Block'
logger = logging.getLogger(__name__)


def add_page_breaks(
        doc: Document,
) -> None:
    """Add page breaks where {{ pagebreak }} markers are found"""
    for para in doc.paragraphs:
        if '{{ pagebreak }}' in para.text:
            para.clear()
            run = para.add_run()
            run.add_break(WD_BREAK.PAGE)


def set_normal_text_style(
        doc: Document,
) -> None:
    """Set the default Normal style"""
    try:
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_after = Pt(6)
    except Exception as e:
        logger.error(f"Failed to set normal text style: {e}")


def create_border_element() -> OxmlElement:
    """Create a border element for paragraphs"""
    border = OxmlElement('w:pBdr')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_elem = OxmlElement(f'w:{border_name}')
        border_elem.set(qn('w:val'), 'single')
        border_elem.set(qn('w:sz'), '4')
        border_elem.set(qn('w:space'), '0')
        border_elem.set(qn('w:color'), 'CCCCCC')
        border.append(border_elem)
    return border


def create_code_block_style(
        doc: Document,
) -> None:
    """Create a custom style for code blocks"""
    try:
        styles = doc.styles
        if CODE_BLOCK not in styles:
            code_style = styles.add_style(CODE_BLOCK, 1)
            code_style.font.name = 'Consolas'
            code_style.font.size = Pt(9)
            code_style.font.color.rgb = RGBColor(0, 0, 0)

            # Paragraph formatting
            pf = code_style.paragraph_format
            pf.left_indent = Inches(0.25)
            pf.right_indent = Inches(0.25)
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)

            # Add background color (light gray)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'F5F5F5')
            code_style._element.get_or_add_pPr().append(shd)  # noqa

            # Add borders
            code_style._element.get_or_add_pPr().append(create_border_element())  # noqa
    except Exception as e:
        logger.error(f"Failed to create code block style: {e}")


def set_cell_borders(
        cell,
) -> None:
    """Set borders for table cells"""
    properties = {
        "top": {"sz": "4", "val": "single", "color": "CCCCCC"},
        "left": {"sz": "4", "val": "single", "color": "CCCCCC"},
        "bottom": {"sz": "4", "val": "single", "color": "CCCCCC"},
        "right": {"sz": "4", "val": "single", "color": "CCCCCC"},
    }

    tc_pr = cell._element.get_or_add_tcPr()  # noqa
    tc_borders = OxmlElement('w:tcBorders')

    for side, attrs in properties.items():
        element = OxmlElement(f'w:{side}')
        for key, value in attrs.items():
            element.set(qn(f'w:{key}'), str(value))
        tc_borders.append(element)

    tc_pr.append(tc_borders)


def set_cell_background(
        cell,
        color: str = 'F9F9F9',
) -> None:
    """Set background color for a cell"""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shd)  # noqa


def set_table_styles(
        doc: Document,
) -> None:
    """Apply custom styling to all tables"""
    for table in doc.tables:
        # Set table width to 100%
        table.autofit = False
        table.allow_autofit = False

        for row in table.rows:
            for cell in row.cells:
                set_cell_borders(cell)
                # Set cell padding
                tc_pr = cell._element.get_or_add_tcPr()  # noqa
                tc_mar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    node = OxmlElement(f'w:{margin_name}')
                    node.set(qn('w:w'), '100')
                    node.set(qn('w:type'), 'dxa')
                    tc_mar.append(node)
                tc_pr.append(tc_mar)

        # Style header row
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                set_cell_background(cell, 'E0E0E0')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)


def apply_heading_styles(
        doc: Document,
) -> None:
    """Apply custom styles to headings"""
    heading_styles = {
        'Heading 1': {
            'font_name': 'Arial',
            'font_size': Pt(20),
            'bold': True,
            'color': RGBColor(31, 78, 121),
            'space_before': Pt(12),
            'space_after': Pt(6),
        },
        'Heading 2': {
            'font_name': 'Arial',
            'font_size': Pt(16),
            'bold': True,
            'color': RGBColor(0, 112, 192),
            'space_before': Pt(10),
            'space_after': Pt(6),
        },
        'Heading 3': {
            'font_name': 'Arial',
            'font_size': Pt(14),
            'bold': True,
            'color': RGBColor(0, 176, 240),
            'space_before': Pt(8),
            'space_after': Pt(4),
        },
        'Heading 4': {
            'font_name': 'Arial',
            'font_size': Pt(12),
            'bold': True,
            'color': RGBColor(0, 0, 0),
            'space_before': Pt(6),
            'space_after': Pt(3),
        },
    }

    for para in doc.paragraphs:
        style_name = para.style.name

        if style_name in heading_styles:
            config = heading_styles[style_name]
            for run in para.runs:
                run.font.name = config['font_name']
                run.font.size = config['font_size']
                run.font.bold = config['bold']
                run.font.color.rgb = config['color']

            para.paragraph_format.space_before = config['space_before']
            para.paragraph_format.space_after = config['space_after']

        # Apply code block style
        if 'code' in style_name.lower() or para.style.name == 'Source Code':
            try:
                para.style = doc.styles[CODE_BLOCK]
            except KeyError:
                pass


def set_page_margins(
        doc: Document,
        margin: float = 0.4,
) -> None:
    """Set page margins for all sections"""
    for section in doc.sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)


def apply_toc_styles(
        doc: Document,
) -> None:
    """Apply consistent styling to Table of Contents"""
    in_toc = False

    for para in doc.paragraphs:
        # Detect start of TOC
        if 'Table of Contents' in para.text:
            in_toc = True
            continue

        # Detect end of TOC
        if in_toc and (
                para.text.strip() == '---' or
                'Detailed Command Reference' in para.text
        ):
            break

        # Apply TOC item styling
        if in_toc and para.text.strip():
            # Reset to Normal style
            para.style = 'Normal'
            para.paragraph_format.outline_level = None

            # Set consistent spacing for ALL TOC items
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.0

            # Fix text formatting for all runs
            for run in para.runs:
                # Normalize size
                run.font.size = Pt(10)

                # Labels should be black, not blue, and NOT bold
                if any(label in run.text for label in ['[Command]', '[Group]', '[CLI]']):
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.underline = False
                    # Explicitly unset bold for [Command] labels
                    run.font.bold = \
                        False if '[Command]' in run.text else run.font.bold


def apply_custom_styles(
        docx_file: str,
) -> None:
    """
    Apply all custom styles to the DOCX document

    Args:
        docx_file: Path to the DOCX file
    """
    try:
        doc = Document(docx_file)

        # Apply various styles
        add_page_breaks(doc)
        set_normal_text_style(doc)
        create_code_block_style(doc)
        apply_heading_styles(doc)
        set_table_styles(doc)
        set_page_margins(doc)

        # Apply TOC styles LAST to ensure consistency
        apply_toc_styles(doc)

        # Save the modified document
        doc.save(docx_file)
        logger.info(f"Successfully applied custom styles to {docx_file}")

    except Exception as e:
        logger.error(f"Failed to apply custom styles: {e}")
        raise
