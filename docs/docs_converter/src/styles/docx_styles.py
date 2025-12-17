"""DOCX styling functions using configuration from style_config."""
import logging
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from .style_config import (
    get_active_style, COLORS, FONTS, SIZES, MARGINS, SPACING, BORDERS,
)

_LOG = logging.getLogger(__name__)
CODE_BLOCK = 'Code Block'

# Heading style mapping: style_name -> (color_key, size_key)
HEADING_STYLE_MAP = {
    'Heading 1': ('heading1', 'heading1'),
    'Heading 2': ('heading2', 'heading2'),
    'Heading 3': ('heading3', 'heading3'),
    'Heading 4': ('heading4', 'heading4'),
    'Heading 5': ('heading5', 'heading4'),
    'Heading 6': ('heading6', 'heading4'),
}


def _create_border_side(
        side: str,
        size: str,
        color: str,
        style: str = 'single'
) -> OxmlElement:
    """
    Create a single border side XML element.

    Args:
        side: Border side ('top', 'left', 'bottom', 'right')
        size: Border size (e.g., '12')
        color: Border color in hex (e.g., '000000')
        style: Border style (e.g., 'single')

    Returns:
        OxmlElement for the border side
    """
    element = OxmlElement(f'w:{side}')
    element.set(qn('w:val'), style)
    element.set(qn('w:sz'), size)
    element.set(qn('w:space'), '0')
    element.set(qn('w:color'), color)
    return element


def set_normal_text_style(doc: Document) -> None:
    """Set the default Normal style."""
    try:
        style_config = get_active_style()
        colors = style_config['colors']
        fonts = style_config['fonts']
        sizes = style_config['sizes']

        style = doc.styles['Normal']
        style.font.name = fonts.get('normal', FONTS['normal'])
        style.font.size = sizes.get('normal_base', SIZES['normal_base'])
        style.font.color.rgb = colors.get('normal', COLORS['normal'])
    except Exception as e:
        _LOG.error(f"Failed to set normal text style: {e}")


def create_border_element() -> OxmlElement:
    """Create border XML element for code blocks."""
    border = OxmlElement('w:pBdr')
    border_color = COLORS['code_border']
    border_size = str(BORDERS['code_size'])
    border_style = BORDERS['border_style']

    for side in ('top', 'left', 'bottom', 'right'):
        border.append(
            _create_border_side(side, border_size, border_color, border_style)
        )

    return border


def create_code_block_style(doc: Document) -> None:
    """Create custom Code Block style."""
    try:
        style_config = get_active_style()
        colors = style_config['colors']
        fonts = style_config['fonts']
        sizes = style_config['sizes']

        styles = doc.styles
        if CODE_BLOCK not in styles:
            code_style = styles.add_style(CODE_BLOCK, WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = fonts.get('code', FONTS['code'])
            code_style.font.size = sizes.get('code', SIZES['code'])
            code_style.font.color.rgb = colors.get('code', COLORS['code'])
            code_style.paragraph_format.left_indent = SPACING['code_left']
            code_style.paragraph_format.right_indent = SPACING['code_right']
            code_style.paragraph_format.space_before = SPACING['code_before']
            code_style.paragraph_format.space_after = SPACING['code_after']

            # Add borders
            p = code_style.paragraph_format
            p_element = p._element
            p_element.get_or_add_pPr().append(create_border_element())
    except Exception as e:
        _LOG.error(f"Failed to create code block style: {e}")


def set_cell_borders(cell) -> None:
    """Set borders for a table cell."""
    border_size = str(BORDERS['table_size'])
    border_color = COLORS['table_border']
    border_style = BORDERS['border_style']

    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')

    for side in ('top', 'left', 'bottom', 'right'):
        tc_borders.append(
            _create_border_side(side, border_size, border_color, border_style)
        )

    tc_pr.append(tc_borders)


def set_table_styles(doc: Document) -> None:
    """Apply custom styling to all tables."""
    for table in doc.tables:
        # Set borders for each cell
        for row in table.rows:
            for cell in row.cells:
                set_cell_borders(cell)

        # Make header row bold
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True


def _apply_paragraph_style(para, doc: Document, style_config: dict) -> None:
    """
    Apply styling to a single paragraph.

    Handles both page breaks and text styling in one pass.
    """
    # Handle page breaks first
    if '{{ pagebreak }}' in para.text:
        para.clear()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)
        return  # Don't process styling for page break paragraphs

    # Get style configuration
    colors = style_config['colors']
    fonts = style_config['fonts']
    sizes = style_config['sizes']

    heading_font = fonts.get('headings', FONTS['headings'])
    normal_font = fonts.get('normal', FONTS['normal'])

    style_name = para.style.name if para.style else 'Normal'

    # Apply styles based on paragraph type
    if style_name == 'Normal':
        for run in para.runs:
            run.font.name = normal_font
            run.font.size = sizes.get('normal', SIZES['normal'])
            run.font.color.rgb = colors.get('normal', COLORS['normal'])

    elif style_name in HEADING_STYLE_MAP:
        color_key, size_key = HEADING_STYLE_MAP[style_name]
        for run in para.runs:
            run.font.name = heading_font
            run.font.size = sizes.get(size_key, SIZES['heading1'])
            run.font.bold = True
            run.font.color.rgb = colors.get(color_key, COLORS['heading1'])

    elif 'code' in style_name.lower():
        try:
            para.style = doc.styles[CODE_BLOCK]
        except KeyError:
            pass


def set_page_margins(doc: Document, margin: float | None = None) -> None:
    """Set page margins for all sections."""
    if margin is None:
        style_config = get_active_style()
        margin = style_config['margins'].get('default', MARGINS['default'])

    for section in doc.sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)


def apply_custom_styles(docx_file: str) -> None:
    """
    Apply all custom styles to a DOCX file.

    Args:
        docx_file: Path to DOCX file to style
    """
    try:
        doc = Document(docx_file)
        style_config = get_active_style()

        # One-time document setup
        _LOG.debug("Setting normal text style...")
        set_normal_text_style(doc)

        _LOG.debug("Creating code block style...")
        create_code_block_style(doc)

        # Single pass through all paragraphs (handles pagebreaks + styling)
        _LOG.debug("Processing paragraphs (page breaks + styling)...")
        for para in doc.paragraphs:
            _apply_paragraph_style(para, doc, style_config)

        # Process tables (separate document structure)
        _LOG.debug("Applying table styles...")
        set_table_styles(doc)

        # Set page margins
        _LOG.debug("Setting page margins...")
        set_page_margins(doc)

        doc.save(docx_file)

        _LOG.info(f"✓ Custom styles applied: {style_config['name']}")

    except Exception as e:
        _LOG.error(f"Failed to apply custom styles: {e}")
        raise
